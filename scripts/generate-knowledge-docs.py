#!/usr/bin/env python3
"""Generate the sample knowledge-base policy documents used by the RAG pipeline.

Writes three real, multi-section policy documents into ``docs/sample-knowledge/``:

    AML-Thresholds-and-Structuring-Policy.docx
    Sanctions-and-High-Risk-Jurisdictions-Policy.pdf
    Cryptocurrency-and-Virtual-Asset-Risk-Policy.docx

The content is deliberately aligned with the thresholds encoded in the seeded
``risk_rules`` (V3__seed.sql), so ``search_policy_knowledge`` returns text that
actually justifies what the deterministic engine computes.

Every section carries a numbered heading (``4.2 Structuring indicators``) so the
section-aware chunker has stable boundaries to split on in both formats: the
.docx headings use real Word ``Heading 1``/``Heading 2`` styles, and the .pdf
headings are emitted as their own bold, larger-font paragraphs on their own line.

Dependencies (install into a throwaway venv if they are not on the system
interpreter)::

    python3 -m venv /tmp/caa-docs-venv
    /tmp/caa-docs-venv/bin/pip install python-docx reportlab pypdf
    /tmp/caa-docs-venv/bin/python scripts/generate-knowledge-docs.py

Usage::

    generate-knowledge-docs.py [--out-dir DIR] [--no-verify]
"""

from __future__ import annotations

import argparse
import sys
from dataclasses import dataclass, field
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_OUT_DIR = REPO_ROOT / "docs" / "sample-knowledge"

INSTALL_HINT = (
    "Missing dependency: {missing}.\n"
    "Install it, for example into a throwaway virtualenv:\n"
    "    python3 -m venv /tmp/caa-docs-venv\n"
    "    /tmp/caa-docs-venv/bin/pip install python-docx reportlab pypdf\n"
    "    /tmp/caa-docs-venv/bin/python scripts/generate-knowledge-docs.py"
)


# ---------------------------------------------------------------------------
# Document model - format independent, rendered to .docx and .pdf below
# ---------------------------------------------------------------------------

@dataclass
class Block:
    kind: str                      # h1 | h2 | p | bullets | table
    text: str = ""
    items: list = field(default_factory=list)
    rows: list = field(default_factory=list)


@dataclass
class Document:
    filename: str
    title: str
    subtitle: str
    reference: str
    blocks: list


def h1(text):
    return Block("h1", text=text)


def h2(text):
    return Block("h2", text=text)


def p(text):
    return Block("p", text=" ".join(text.split()))


def bullets(*items):
    return Block("bullets", items=[" ".join(i.split()) for i in items])


def table(header, *rows):
    return Block("table", rows=[list(header)] + [list(r) for r in rows])


# ---------------------------------------------------------------------------
# Document 1 - AML thresholds and structuring
# ---------------------------------------------------------------------------

AML_POLICY = Document(
    filename="AML-Thresholds-and-Structuring-Policy.docx",
    title="Anti-Money-Laundering Transaction Monitoring Policy",
    subtitle="Reporting thresholds, structuring detection and escalation duties",
    reference="Reference AML-POL-004 | Version 6.2 | Owner: Group Financial Crime Compliance",
    blocks=[
        h1("1. Purpose and scope"),
        p("""This policy defines the monetary thresholds, behavioural indicators and escalation
           duties that govern automated transaction monitoring across all customer payment,
           card and virtual-asset activity. It applies to every booking entity and to every
           analyst, whether human or automated, that produces a risk assessment on a customer
           file. Nothing in this policy replaces the analyst's duty to escalate a suspicion that
           no rule has captured."""),
        p("""Transaction monitoring in this institution is deliberately asymmetric. The cost of a
           false positive is one additional review; the cost of a false negative is an
           undetected predicate offence, a regulatory finding and reputational damage. Where the
           evidence is ambiguous, the case is escalated, not cleared."""),

        h1("2. Reporting thresholds"),
        p("""Two thresholds govern payment monitoring: the single-instruction reporting threshold
           of 10,000 currency units, and the aggregated 24-hour threshold of 20,000 currency
           units across three or more instructions. Both are absolute; neither is subject to
           analyst discretion, and neither is waived by the customer's tenure or segment."""),
        h2("2.1 The 10,000 currency-unit reporting threshold"),
        p("""Any single payment instruction with a settled value at or above 10,000 currency units
           is a reportable transaction. Reportable transactions are recorded, retained for ten
           years and, where the ordering or beneficiary party cannot be fully identified,
           reported to the national Financial Intelligence Unit. The threshold applies to the
           settled amount, not to the instructed amount, and is not reduced by fees."""),
        p("""The threshold is currency-unit based and is applied without conversion tolerance. An
           instruction of 9,999.99 is not reportable on its face; an instruction of 10,000.00 is.
           That cliff edge is precisely what makes the band immediately below it the single most
           informative range in the whole data set."""),
        table(
            ("Band", "Treatment", "Standard action"),
            ("Below 3,000", "Routine", "No monitoring action unless a pattern rule fires"),
            ("3,000 to 7,999.99", "Monitored", "Aggregated into velocity and volume statistics"),
            ("8,000 to 9,999.99", "Threshold-adjacent", "Structuring screen, see section 4"),
            ("10,000 and above", "Reportable", "Record, retain, review the beneficiary"),
            ("75,000 and above", "Enhanced review", "Second-line sign-off before release"),
        ),
        h2("2.2 Aggregation windows"),
        p("""A single instruction below the threshold does not exempt a customer from reporting
           duties. Instructions are aggregated over a rolling 24-hour window and a rolling
           30-day window per customer. Where the aggregate of a rolling 24-hour window reaches
           20,000 currency units across three or more instructions, the window itself is treated
           as a reportable event even though no individual instruction crossed the threshold."""),
        p("""Aggregation is performed per customer, not per account. Customers who hold several
           accounts or several payment instruments are aggregated at the customer level, because
           splitting a payment across instruments is itself a recognised evasion technique."""),

        h1("3. Velocity and value spikes"),
        p("""A velocity spike is defined as eight or more transactions of any activity type inside
           a rolling 24-hour window whose aggregate value exceeds 40,000 currency units. Velocity
           spikes are not inherently suspicious - salary runs, supplier settlement days and
           property completions all produce them - but they must be explained. An unexplained
           velocity spike on a retail customer profile is a strong indicator that the account is
           being used as a pass-through."""),
        bullets(
            "Compare the spike against the customer's own 30-day baseline, not against a portfolio average.",
            "Check whether the spike is composed of round-number amounts, which suggests instruction rather than commerce.",
            "Check whether funds leave the account within 48 hours of arriving, which is the classic funnel-account signature.",
            "Record the explanation. An unexplained spike that recurs in a later period is an escalation, not a repeat review.",
        ),
        p("""High-value activity booked outside normal business hours deserves separate attention.
           Instructions above 15,000 currency units executed between 00:00 and 05:59 local
           booking time are flagged, because legitimate corporate and retail flows overwhelmingly
           settle during business hours while instructed fraud and coerced payments cluster at
           night."""),

        h1("4. Structuring and smurfing"),
        p("""Structuring is the pattern this policy treats most seriously among purely monetary
           indicators, because unlike a large payment it has no innocent mechanical explanation:
           the amounts are chosen by the customer with reference to a threshold the customer
           should not be optimising against. This section defines the offence, lists the
           indicators the engine screens for, and sets out how to separate it from the legitimate
           splitting of a payment."""),
        h2("4.1 Definition"),
        p("""Structuring is the deliberate division of a larger sum into a series of smaller
           instructions, each individually below a reporting threshold, with the purpose of
           avoiding the reporting duty. It is a criminal offence in its own right in most
           jurisdictions in which this institution operates, independently of whether the
           underlying funds are of criminal origin. Structuring is therefore reportable even
           where the source of funds is later shown to be legitimate."""),
        h2("4.2 Structuring indicators"),
        p("""No single indicator is conclusive. The following pattern, however, is treated as a
           presumptive structuring case and requires a Suspicious Activity Report unless the
           customer file already carries a documented, dated and plausible explanation."""),
        bullets(
            "Three or more payments in the 8,000 to 9,999.99 band inside a rolling 24-hour window.",
            "An aggregate of 20,000 currency units or more across that same window.",
            "Amounts that cluster just below the threshold rather than being distributed across the plausible range.",
            "Use of more than one payment rail - for example ACH and peer-to-peer on the same day - for economically identical transfers.",
            "Repetition of the pattern in separate bursts several weeks apart, which distinguishes structuring from a one-off large purchase that was split for operational reasons.",
        ),
        h2("4.3 Distinguishing structuring from legitimate splitting"),
        p("""Payments are legitimately split when a counterparty imposes a per-instruction limit,
           when an escrow schedule requires staged release, or when a card scheme or rail imposes
           its own cap. In each of those cases the split is externally imposed and the amounts
           follow the external limit, not the reporting threshold. The distinguishing question is
           always: does the amount track a commercial constraint, or does it track 10,000?"""),
        p("""A customer whose instructions repeatedly land between 9,200 and 9,975 is tracking the
           threshold. A customer whose instructions land at exactly 5,000 because that is the
           rail's per-instruction cap is tracking the rail. The first is escalated; the second is
           documented and cleared."""),

        h1("5. Escalation and reporting duties"),
        p("""Where the monitoring engine produces a total risk score in the HIGH or CRITICAL band,
           a second-line reviewer must be assigned within one business day. Where a structuring
           or sanctions rule has fired, the review must be completed and the reporting decision
           documented within three business days."""),
        table(
            ("Score band", "Risk level", "Required action"),
            ("0 to 24", "LOW", "Automated file note, no manual review"),
            ("25 to 49", "MEDIUM", "First-line review within five business days"),
            ("50 to 74", "HIGH", "Second-line review within one business day"),
            ("75 and above", "CRITICAL", "Immediate second-line review, consider account restriction"),
        ),
        p("""Every rule in the monitoring rule set must be evaluated on every analysis, including
           rules that do not fire. A rule that was not evaluated is treated for audit purposes as
           a rule that failed, because an unevaluated rule cannot support the conclusion that no
           risk was present."""),
    ],
)


# ---------------------------------------------------------------------------
# Document 2 - Sanctions and high-risk jurisdictions
# ---------------------------------------------------------------------------

SANCTIONS_POLICY = Document(
    filename="Sanctions-and-High-Risk-Jurisdictions-Policy.pdf",
    title="Sanctions and High-Risk Jurisdictions Policy",
    subtitle="Beneficiary screening, prohibited destinations and correspondent banking controls",
    reference="Reference SAN-POL-002 | Version 9.0 | Owner: Group Sanctions Compliance",
    blocks=[
        h1("1. Purpose and legal basis"),
        p("""This policy sets out how cross-border payment instructions are screened against
           sanctions regimes and how jurisdictional risk is scored. It implements the obligations
           arising from the applicable United Nations, European Union, United States OFAC, United
           Kingdom OFSI and Swiss SECO measures, and the Financial Action Task Force public
           statements on jurisdictions under increased monitoring."""),
        p("""Screening is applied to the beneficiary bank country, the beneficiary account, the
           ordering customer and any intermediary institution named in the instruction. A hit on
           any one of those four is sufficient to hold the instruction."""),

        h1("2. Prohibited and restricted jurisdictions"),
        p("""Jurisdictional risk is expressed as two lists. The first, in section 2.1, carries
           comprehensive or sectoral measures and blocks the instruction outright. The second, in
           section 3, carries no sanctions but demands beneficial-ownership evidence above a
           monetary floor. An instruction is screened against both lists on the beneficiary bank
           country recorded on the payment."""),
        h2("2.1 Comprehensive measures"),
        p("""The jurisdictions below are subject to comprehensive or near-comprehensive measures.
           A payment whose beneficiary bank is located in any of them is blocked pending
           second-line sanctions review, regardless of amount, and is never released on a
           first-line decision."""),
        table(
            ("ISO-2", "Jurisdiction", "Regime"),
            ("IR", "Iran", "Comprehensive - UN, EU, OFAC, SECO"),
            ("KP", "North Korea", "Comprehensive - UN, EU, OFAC, SECO"),
            ("SY", "Syria", "Comprehensive - EU, OFAC, SECO"),
            ("CU", "Cuba", "Comprehensive - OFAC"),
            ("RU", "Russian Federation", "Sectoral and financial - EU, OFAC, OFSI, SECO"),
            ("BY", "Belarus", "Sectoral and financial - EU, OFAC, SECO"),
            ("AF", "Afghanistan", "Targeted - UN, EU, OFAC"),
            ("MM", "Myanmar", "Targeted - EU, OFAC"),
            ("VE", "Venezuela", "Targeted, government sector - OFAC"),
        ),
        h2("2.2 Consequences of a hit"),
        p("""A beneficiary bank country appearing in the table above is by itself a sufficient
           basis for escalation. The instruction is not a permitted exception merely because the
           beneficiary is a private individual, because the amount is small, or because the
           customer has transacted with that jurisdiction before without a report being filed.
           Prior unreported activity to a listed jurisdiction is itself a look-back trigger."""),
        p("""Payments routed through a correspondent in a listed jurisdiction are treated
           identically to payments whose ultimate beneficiary is in that jurisdiction. Routing
           through an unlisted intermediary does not cleanse the instruction."""),

        h1("3. Offshore and secrecy jurisdictions"),
        p("""The following jurisdictions are not subject to sanctions but present elevated
           beneficial-ownership opacity: Panama, the Cayman Islands, the Seychelles, the British
           Virgin Islands, the Bahamas and Belize. Activity to these destinations is permitted,
           but a single instruction above 25,000 currency units to a beneficiary bank in one of
           them requires documented beneficial-ownership evidence on the customer file before
           release."""),
        bullets(
            "Verify the beneficiary is not a shell entity incorporated within the previous twelve months.",
            "Obtain the ultimate beneficial owner, not merely the registered agent.",
            "Where the beneficiary is a trust or foundation, obtain the settlor and the class of beneficiaries.",
            "Record the commercial rationale. 'Investment' without a named instrument is not a rationale.",
        ),

        h1("4. Wire and correspondent banking controls"),
        p("""Wire and SWIFT instructions carry the highest per-instruction value in the book and
           the least visibility into the onward chain. Two controls apply on top of list
           screening: an amount-based enhanced review on individual instructions, and a
           pattern-based control on the geographic spread of a customer's instructions over
           time."""),
        h2("4.1 High-value SWIFT instructions"),
        p("""A SWIFT instruction above 75,000 currency units to a beneficiary bank outside the
           domestic booking jurisdiction is subject to enhanced review before release. The
           reviewer confirms the beneficiary bank's regulatory status, the presence of complete
           originator information under the FATF travel rule, and the consistency of the payment
           with the customer's declared business."""),
        p("""Incomplete originator or beneficiary information is a standalone escalation. An
           instruction that omits the ordering customer's address, or that names a beneficiary
           only as an account number, must not be released on the basis that the correspondent
           will complete it downstream."""),
        h2("4.2 Cross-border fan-out"),
        p("""Fan-out is the pattern in which one customer sends wire or SWIFT instructions above
           10,000 currency units to beneficiary banks in five or more distinct countries inside a
           rolling 30-day window. Fan-out is characteristic of layering: the purpose is to break
           the audit trail by dispersing value across jurisdictions and correspondent
           relationships faster than any single investigator can follow it."""),
        bullets(
            "Establish whether the customer's declared business plausibly requires that geographic spread.",
            "Check for repeated use of the same beneficiary name across several jurisdictions.",
            "Check whether any leg of the fan-out terminates in a listed or offshore jurisdiction.",
            "Where fan-out coincides with a sanctioned-jurisdiction leg, treat the whole pattern as one case, not as separate alerts.",
        ),

        h1("5. Customer domicile and residual risk"),
        p("""Customer domicile is a contributing factor, never a conclusion. A customer resident in
           a higher-risk jurisdiction who transacts domestically and transparently presents lower
           residual risk than a customer resident in a low-risk jurisdiction who wires value into
           a comprehensively sanctioned one. Analysts must score the flow, not the passport."""),
        p("""Where a customer's activity has triggered a sanctions rule, the analyst must record
           the specific instruction identifiers that caused it. A narrative that asserts exposure
           to a sanctioned jurisdiction without naming the underlying transactions is not an
           acceptable assessment and will be returned by quality assurance."""),
    ],
)


# ---------------------------------------------------------------------------
# Document 3 - Cryptocurrency and virtual asset risk
# ---------------------------------------------------------------------------

CRYPTO_POLICY = Document(
    filename="Cryptocurrency-and-Virtual-Asset-Risk-Policy.docx",
    title="Cryptocurrency and Virtual Asset Risk Policy",
    subtitle="Privacy chains, mixing services, counterparty attribution and exposure limits",
    reference="Reference VAS-POL-007 | Version 4.1 | Owner: Digital Assets Financial Crime",
    blocks=[
        h1("1. Purpose and scope"),
        p("""This policy governs the monitoring of customer virtual-asset activity, including
           transfers on public blockchains, transfers to and from virtual-asset service providers
           and any activity settled in a stablecoin. It applies to the on-chain leg of a transfer
           as well as to the fiat leg."""),
        p("""Virtual-asset activity is not treated as inherently suspicious. It is treated as
           activity whose counterparty is unknown until proved otherwise, which reverses the
           default assumption that applies to correspondent banking."""),

        h1("2. Counterparty attribution"),
        p("""Attribution is the single most informative attribute of a virtual-asset transfer.
           A transfer of any size to a named, regulated exchange is ordinary banking; the same
           transfer to an address that cannot be attributed to any known service is an
           unresolved counterparty and is scored as such."""),
        h2("2.1 The attribution requirement"),
        p("""Every virtual-asset transfer must carry an attributed counterparty: a named,
           regulated virtual-asset service provider, a known merchant processor, or a wallet the
           customer has demonstrably self-declared. A transfer whose counterparty exchange cannot
           be attributed is an unattributed transfer."""),
        p("""An empty exchange attribution is a substantive finding, not a data-quality issue. It
           means the destination is either a self-hosted wallet, a service outside the analytics
           provider's coverage, or a service that deliberately resists attribution. All three
           materially raise the risk of the transfer."""),
        h2("2.2 Escalation on unattributed transfers"),
        bullets(
            "A single unattributed transfer to a self-hosted wallet is documented but not escalated on its own.",
            "Repeated unattributed transfers to the same destination address are escalated.",
            "An unattributed transfer that is also a privacy-chain or mixer transfer is escalated immediately, without an amount floor.",
            "Unattributed inbound value that is forwarded onward within 24 hours is treated as pass-through and escalated.",
        ),

        h1("3. Privacy chains and anonymity-enhancing technology"),
        p("""Anonymity-enhancing technology comes in two forms that this policy treats
           differently: chains whose protocol conceals the transfer, covered in section 3.1, and
           services layered on a transparent chain that deliberately break the link between
           source and destination, covered in section 3.2. The first is a risk finding; the
           second, where the service is designated, is a sanctions finding."""),
        h2("3.1 Designated privacy chains"),
        p("""Monero (XMR), Zcash (ZEC) shielded transactions and Dash (DASH) PrivateSend are
           designated anonymity-enhanced. These protocols conceal amounts, addresses or both by
           design, which makes downstream tracing impossible rather than merely difficult. The
           institution does not accept the position that a transfer on such a chain can be
           risk-assessed from on-chain data alone."""),
        p("""A transfer on a designated privacy chain that also lacks an attributed exchange is
           the highest-severity virtual-asset finding in this policy. There is no legitimate
           retail pattern in which value repeatedly leaves a customer's control on an
           anonymity-enhanced chain toward a counterparty that cannot be named."""),
        h2("3.2 Mixing and tumbling services"),
        p("""Mixing services, tumblers and privacy pools pool deposits from many users and return
           value from an unrelated pool, breaking the deterministic on-chain link between source
           and destination. Several such services are designated by the United States Office of
           Foreign Assets Control, including Tornado Cash and Blender.io together with its
           successor Sinbad.io. Transfers to designated service addresses are sanctions matters,
           not merely risk-scoring matters."""),
        bullets(
            "Screen every destination address against the designated-address list before release.",
            "Treat a hit on a designated mixer address as a sanctions escalation with the same urgency as a comprehensively sanctioned jurisdiction.",
            "Record the specific transaction hashes and destination addresses in the assessment; a narrative without addresses is not evidence.",
            "Do not net mixer exposure against unrelated legitimate activity. Exposure is assessed gross.",
        ),
        p("""The addresses below are the designated service addresses currently screened by the
           monitoring engine. The list is maintained by Group Sanctions Compliance and is not
           exhaustive; an address absent from it is unscreened, not cleared."""),
        table(
            ("Destination address", "Designated service"),
            ("0x8589427373D6D84E98730D7795D8f6f8731FDA16", "Tornado Cash - OFAC SDN"),
            ("0x722122dF12D4e14e13Ac3b6895a86e84145b6967", "Tornado Cash router - OFAC SDN"),
            ("0xDD4c48C0B24039969fC16D1cdF626eaB821d3384", "Tornado Cash 0.1 ETH pool - OFAC SDN"),
            ("bc1qm34lsc65zpw79lxes69zkqmk6ee3ewf0j77s3h", "Blender.io / Sinbad.io - OFAC SDN"),
        ),

        h1("4. Concentration and exposure limits"),
        p("""Where virtual-asset transactions represent 70 percent or more of a customer's total
           transaction count over a rolling 30-day window, the customer is treated as a
           virtual-asset-concentrated profile. Concentration is not itself an adverse finding,
           but it changes the baseline: subsequent single transfers above 5,000 currency units on
           a concentrated profile are reviewed rather than sampled."""),
        table(
            ("Indicator", "Threshold", "Treatment"),
            ("30-day crypto share of activity", "70 percent or more", "Concentrated profile"),
            ("Single transfer on a concentrated profile", "Above 5,000", "Individual review"),
            ("Privacy-chain transfer, unattributed", "Any amount", "Immediate escalation"),
            ("Designated mixer address", "Any amount", "Sanctions escalation"),
            ("Onward forwarding of inbound value", "Within 24 hours", "Pass-through review"),
        ),
        p("""Fiat on-ramp and off-ramp activity is assessed against the same monetary thresholds as
           any other payment. A purchase of virtual assets funded by a series of card
           authorisations just below a per-transaction limit is a structuring pattern and is
           handled under the anti-money-laundering policy, not under this one."""),

        h1("5. Card-based virtual asset purchases"),
        p("""Card authorisations at quasi-cash and crypto-kiosk merchant category codes,
           in particular MCC 6051, are frequently used to fund virtual-asset purchases with
           compromised card credentials. A burst of declined authorisations at such a merchant
           followed by a successful card-not-present authorisation is a recognised card-testing
           and cash-out sequence."""),
        bullets(
            "Five or more declined authorisations by the same customer inside 24 hours is a velocity finding in its own right.",
            "A card-not-present success above 3,000 currency units following four or more declines in the same window is treated as a probable compromise.",
            "Related merchant category codes carry elevated risk: 7995 gambling, 6051 quasi-cash, 4829 money transfer, 6211 securities, 7273 dating services.",
            "Freeze the instrument before contacting the customer where the sequence is present; contacting first gives a fraudster time to complete the cash-out.",
        ),

        h1("6. Evidence standards for virtual-asset assessments"),
        p("""An assessment of virtual-asset risk must name the blockchain, the destination address,
           the transaction hash and the attributed counterparty or its absence for every transfer
           relied upon. Aggregate statements such as 'significant crypto exposure' are not
           acceptable without the underlying transaction identifiers, because they cannot be
           reviewed, reproduced or defended to a regulator."""),
        p("""Where an automated engine and a human or model-produced narrative disagree about
           whether a rule fired, the deterministic engine result governs the score and the
           disagreement is recorded in the case file. The narrative may argue for a higher risk
           level than the score implies; it may never argue for a lower one without new evidence
           attached to the file."""),
    ],
)


DOCUMENTS = [AML_POLICY, SANCTIONS_POLICY, CRYPTO_POLICY]


# ---------------------------------------------------------------------------
# Renderers
# ---------------------------------------------------------------------------

def render_docx(doc: Document, path: Path) -> None:
    from docx import Document as DocxDocument
    from docx.shared import Pt

    d = DocxDocument()
    d.core_properties.title = doc.title
    d.core_properties.author = "Group Financial Crime Compliance"
    d.core_properties.category = "Policy"

    d.add_heading(doc.title, level=0)
    sub = d.add_paragraph(doc.subtitle)
    sub.runs[0].italic = True
    ref = d.add_paragraph(doc.reference)
    ref.runs[0].font.size = Pt(9)

    for block in doc.blocks:
        if block.kind == "h1":
            d.add_heading(block.text, level=1)
        elif block.kind == "h2":
            d.add_heading(block.text, level=2)
        elif block.kind == "p":
            d.add_paragraph(block.text)
        elif block.kind == "bullets":
            for item in block.items:
                d.add_paragraph(item, style="List Bullet")
        elif block.kind == "table":
            t = d.add_table(rows=len(block.rows), cols=len(block.rows[0]))
            t.style = "Light Grid Accent 1"
            for r, row in enumerate(block.rows):
                for c, cell in enumerate(row):
                    t.cell(r, c).text = str(cell)
            for cell in t.rows[0].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
            d.add_paragraph()
    d.save(str(path))


def render_pdf(doc: Document, path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (ListFlowable, ListItem, Paragraph,
                                    SimpleDocTemplate, Spacer, Table, TableStyle)

    base = getSampleStyleSheet()
    title_style = ParagraphStyle("DocTitle", parent=base["Title"], fontSize=19, leading=23,
                                 spaceAfter=4)
    subtitle_style = ParagraphStyle("DocSubtitle", parent=base["Normal"], fontSize=11,
                                    leading=14, textColor=colors.HexColor("#444444"),
                                    fontName="Helvetica-Oblique", spaceAfter=2)
    ref_style = ParagraphStyle("DocRef", parent=base["Normal"], fontSize=8, leading=11,
                               textColor=colors.HexColor("#666666"), spaceAfter=14)
    # Headings are their own bold, larger-font paragraph on their own line, so a
    # section-aware chunker can recover them from the extracted text stream.
    h1_style = ParagraphStyle("H1", parent=base["Heading1"], fontSize=14.5, leading=18,
                              spaceBefore=16, spaceAfter=7,
                              textColor=colors.HexColor("#0d1b2a"))
    h2_style = ParagraphStyle("H2", parent=base["Heading2"], fontSize=12, leading=15,
                              spaceBefore=11, spaceAfter=5,
                              textColor=colors.HexColor("#1b3a5c"))
    body_style = ParagraphStyle("Body", parent=base["BodyText"], fontSize=10, leading=14.5,
                                alignment=TA_JUSTIFY, spaceAfter=7)
    bullet_style = ParagraphStyle("Bullet", parent=body_style, spaceAfter=3)

    story = [Paragraph(doc.title, title_style),
             Paragraph(doc.subtitle, subtitle_style),
             Paragraph(doc.reference, ref_style)]

    for block in doc.blocks:
        if block.kind == "h1":
            story.append(Paragraph(block.text, h1_style))
        elif block.kind == "h2":
            story.append(Paragraph(block.text, h2_style))
        elif block.kind == "p":
            story.append(Paragraph(block.text, body_style))
        elif block.kind == "bullets":
            story.append(ListFlowable(
                [ListItem(Paragraph(item, bullet_style), leftIndent=12) for item in block.items],
                bulletType="bullet", start="-", leftIndent=14, bulletFontSize=9))
            story.append(Spacer(1, 7))
        elif block.kind == "table":
            widths = {2: [70 * mm, 96 * mm], 3: [38 * mm, 52 * mm, 76 * mm]}
            t = Table(block.rows, hAlign="LEFT",
                      colWidths=widths.get(len(block.rows[0])))
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0d1b2a")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#b0b8c1")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                 [colors.white, colors.HexColor("#f2f4f7")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.append(t)
            story.append(Spacer(1, 10))

    template = SimpleDocTemplate(
        str(path), pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm, topMargin=18 * mm, bottomMargin=18 * mm,
        title=doc.title, author="Group Sanctions Compliance", subject=doc.subtitle)
    template.build(story)


# ---------------------------------------------------------------------------
# Verification - the files must actually open and yield extractable text
# ---------------------------------------------------------------------------

def verify_docx(path: Path) -> str:
    from docx import Document as DocxDocument

    d = DocxDocument(str(path))
    headings = [para.text for para in d.paragraphs if para.style.name.startswith("Heading")]
    words = sum(len(para.text.split()) for para in d.paragraphs)
    if not headings:
        raise SystemExit(f"{path.name}: no Word heading styles found")
    if words < 400:
        raise SystemExit(f"{path.name}: only {words} words of body text")
    return f"{len(d.paragraphs)} paragraphs, {len(headings)} headings, {len(d.tables)} tables, {words} words"


def verify_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        header = path.open("rb").read(5)
        if header != b"%PDF-":
            raise SystemExit(f"{path.name}: not a PDF")
        return "PDF header valid (install pypdf for full text-extraction check)"

    reader = PdfReader(str(path))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    numbered = [line.strip() for line in text.splitlines()
                if line.strip()[:1].isdigit() and "." in line.strip()[:4]]
    words = len(text.split())
    if words < 400:
        raise SystemExit(f"{path.name}: only {words} words extracted")
    if len(numbered) < 5:
        raise SystemExit(f"{path.name}: only {len(numbered)} numbered headings recoverable")
    return f"{len(reader.pages)} pages, {words} extractable words, {len(numbered)} numbered headings"


# ---------------------------------------------------------------------------

def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__,
                                     formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--out-dir", type=Path, default=DEFAULT_OUT_DIR,
                        help=f"output directory (default: {DEFAULT_OUT_DIR})")
    parser.add_argument("--no-verify", action="store_true",
                        help="skip re-opening the generated files")
    args = parser.parse_args()

    missing = []
    try:
        import docx  # noqa: F401
    except ImportError:
        missing.append("python-docx")
    try:
        import reportlab  # noqa: F401
    except ImportError:
        missing.append("reportlab")
    if missing:
        print(INSTALL_HINT.format(missing=", ".join(missing)), file=sys.stderr)
        return 2

    out_dir = args.out_dir
    out_dir.mkdir(parents=True, exist_ok=True)

    for doc in DOCUMENTS:
        path = out_dir / doc.filename
        if path.suffix == ".docx":
            render_docx(doc, path)
        elif path.suffix == ".pdf":
            render_pdf(doc, path)
        else:
            raise SystemExit(f"unsupported output format: {path.suffix}")
        size = path.stat().st_size
        detail = ""
        if not args.no_verify:
            detail = "  ->  " + (verify_docx(path) if path.suffix == ".docx" else verify_pdf(path))
        print(f"{path.relative_to(REPO_ROOT) if REPO_ROOT in path.parents else path}"
              f"  ({size:,} bytes){detail}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
